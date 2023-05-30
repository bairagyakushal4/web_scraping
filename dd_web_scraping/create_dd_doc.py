import io
from PIL import Image
from bs4 import BeautifulSoup
import requests
import openpyxl
from docx import Document
from docx.shared import Cm, Pt


# 1 px = ( 2.54 / 96 ) cm
# width_in_cm = width_in_px * (2.54 / 96)
# height_in_cm = height_in_px * (2.54 / 96)
# width_in_cm = float("%.2f" % width_in_cm)
# height_in_cm = float("%.2f" % height_in_cm)
# print('cm', width_in_cm, height_in_cm)


reduce_page_break_works = 1.2

# for img add to docx
max_width = 17.79
max_height = 22.86 - reduce_page_break_works


def init_document(doc):
    # core properties
    core_properties = doc.core_properties
    core_properties.title = 'Dungeon Defense (WN)'
    core_properties.author = 'Yoo Heonhwa'
    core_properties.comments = 'Generated with python by Kushal'

    # style the body text
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(16)

    # style the heading
    HeadingStyle = doc.styles['Heading 1']
    HeadingFont = HeadingStyle.font
    HeadingFont.size = Pt(24)

    # set margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)


def Scrape_Site(site_url, chapter_heading):
    try:
        source = requests.get(site_url)
        source.raise_for_status()
        soup = BeautifulSoup(source.text, 'html.parser')

        storyPart = soup.select_one(".entry-content")

        # remover unnecessary tags
        for header in storyPart.select('style'):
            header.decompose()

        for header in storyPart.select('noscript'):
            header.decompose()

        for header in storyPart.select('script'):
            header.decompose()

        doc.add_heading(chapter_heading, level=1)
        finalStory = storyPart.findChildren()

        # skip_tag is used to iterate over all the child of a tag and not get duplicate content of findChildren method
        skip_tag = None

        exclude_text = [
            'PREVIOUS CHAPTER | NEXT CHAPTER',
            'NEXT CHAPTER',
            'PREVIOUS CHAPTER'
        ]

        for t, tag in enumerate(finalStory):

            # no of tag in the chapter
            # print(t)
            # if t>5

            # check if img exist in the paragraph add text
            imgs = tag.findAll("img")
            if imgs != []:
                for img in imgs:
                    doc.add_page_break()
                    img_url = img['src']
                    img_response = requests.get(img_url, stream=True)
                    image = io.BytesIO(img_response.content)
                    # get height and width of the img and adjust inside the page
                    img = Image.open(image)
                    width_in_px, height_in_px = img.size

                    if width_in_px > height_in_px:
                        width_in_cm = max_width
                        doc.add_picture(image, width=Cm(width_in_cm))

                    elif height_in_px > width_in_px:
                        height_in_px = max_height
                        doc.add_picture(image, height=Cm(height_in_px))

                    doc.add_page_break()

            else:

                break_flag = 0
                continue_flag = 0

                current_tag_text = tag.get_text(strip=True)

                # exclude black nad new line
                if len(current_tag_text) <= 1:
                    continue_flag = continue_flag + 1
                    # print('====blank====')

                for ex in exclude_text:
                    if ex in current_tag_text:
                        if t > 5:
                            break_flag = break_flag + 1
                            break

                if continue_flag > 0:
                    continue

                # no need to continue after previous/next chapter anchor tag found
                if break_flag > 0:
                    break

                storyPart_text = tag.get_text("\n\n", strip=True)
                storyPart_text = storyPart_text.replace("Ο", "\n")
                # storyPart_text = storyPart_text.replace("\n\n\n", "\n")

                # loop through the tag and it's all children
                if skip_tag != tag:
                    doc.add_paragraph('\n' + storyPart_text)

                    tagChildren = tag.findChildren()
                    if tagChildren != []:
                        # skip the first tag Child during next loop
                        skip_tag = tagChildren[0]
                    else:
                        skip_tag = None
        # chapter complete start from next page
        doc.add_page_break()

    except Exception as e:
        print(e)


doc = Document()
init_document(doc)


excel = openpyxl.load_workbook('Dungeon-Defense-(WN)-Table-of-Contents.xlsx')
sheet = excel.active

total_rows = sheet.max_row

all_ch_title = tuple()
all_site_url = tuple()


for x in range(1, total_rows + 1):
    col_title = sheet.cell(row=x, column=1).value
    col_link = sheet.cell(row=x, column=2).value

    all_ch_title += (col_title,)
    all_site_url += (col_link,)


count = i = 0
countGallop = countGallopGap = 50  # no of ch per file

doc = Document()
init_document(doc)


fileNameMain = "Dungeon-Defense-(WN)-{}.docx"

for site_url in all_site_url:
    count = count + 1
    chapter_heading = all_ch_title[count-1]

    Scrape_Site(site_url, chapter_heading)

    # write a specific no of chapter in multiple files
    if count == countGallop:
        i = i + 1
        countGallop = countGallop + countGallopGap

        doc.save(fileNameMain.format(i))

        doc = Document()
        init_document(doc)
        print(count, 'if', chapter_heading)

    # save the last remaining no of chapter in another file
    elif count == len(all_site_url):
        i = i + 1
        doc.save(fileNameMain.format(i))

    else:
        print(count, 'else', chapter_heading)

    # # testing --comment below 2 line to get the desired output
    # if count == countGallopGap:
    #     break
