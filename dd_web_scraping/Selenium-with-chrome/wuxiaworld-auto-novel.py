from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By

from selenium.webdriver.chrome.options import Options
import time

import io
from PIL import Image
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Cm, Pt


reduce_page_break_works = 1.2

# for img add to docx
max_width = 17.79
max_height = 22.86 - reduce_page_break_works


def init_document(doc):
    # core properties
    core_properties = doc.core_properties
    core_properties.title = 'Apartment for Rent'
    core_properties.author = 'Giddens Ko'
    core_properties.comments = 'Generated with python by Kushal'

    # style the body text
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(16)

    # style the heading
    HeadingStyle = doc.styles['Heading 1']
    HeadingFont = HeadingStyle.font
    HeadingFont.size = Pt(24)

    # set margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)


def Scrape_Site(content, chapter_heading):
    try:
        source = content.get_attribute('innerHTML')
        soup = BeautifulSoup(source, 'html.parser')

        storyPart = soup

        # remover unnecessary tags
        for header in storyPart.select('style'):
            header.decompose()

        for header in storyPart.select('noscript'):
            header.decompose()

        for header in storyPart.select('script'):
            header.decompose()

        for header in storyPart.select('button'):
            header.decompose()

        doc.add_heading(chapter_heading, level=1)
        finalStory = storyPart.findChildren()

        # skip_tag is used to iterate over all the child of a tag and not get duplicate content of findChildren method
        skip_tag = None

        exclude_text = [
            'Previous Chapter | Next Chapter',
            'Next Chapter',
            'Previous Chapter'
        ]

        for t, tag in enumerate(finalStory):

            # no of tag in the chapter
            # print(t)
            # if t>5

            # check if img exist in the paragraph add text
            imgs = tag.findAll("img")
            if imgs != []:
                for img in imgs:
                    doc.add_page_break()
                    img_url = img['src']
                    img_response = requests.get(img_url, stream=True)
                    image = io.BytesIO(img_response.content)
                    # get height and width of the img and adjust inside the page
                    img = Image.open(image)
                    width_in_px, height_in_px = img.size

                    if width_in_px > height_in_px:
                        width_in_cm = max_width
                        doc.add_picture(image, width=Cm(width_in_cm))

                    elif height_in_px > width_in_px:
                        height_in_px = max_height
                        doc.add_picture(image, height=Cm(height_in_px))

                    doc.add_page_break()

            else:

                break_flag = 0
                continue_flag = 0

                current_tag_text = tag.get_text(strip=True)

                # exclude black nad new line
                if len(current_tag_text) <= 1:
                    continue_flag = continue_flag + 1
                    # print('====blank====')

                for ex in exclude_text:
                    if ex in current_tag_text:
                        if t > 5:
                            break_flag = break_flag + 1
                            break

                if continue_flag > 0:
                    continue

                # no need to continue after previous/next chapter anchor tag found
                if break_flag > 0:
                    break

                storyPart_text = tag.get_text("\n\n", strip=True)
                storyPart_text = storyPart_text.replace("Ο", "\n")
                # storyPart_text = storyPart_text.replace("\n\n\n", "\n")

                # loop through the tag and it's all children
                if skip_tag != tag:
                    doc.add_paragraph('\n' + storyPart_text)

                    tagChildren = tag.findChildren()
                    if tagChildren != []:
                        # skip the first tag Child during next loop
                        skip_tag = tagChildren[0]
                    else:
                        skip_tag = None
        # chapter complete start from next page
        doc.add_page_break()

    except Exception as e:
        print(e)


doc = Document()
init_document(doc)


fileNameMain = "Apartment-for-Rent.docx"


service = Service(executable_path="chromedriver.exe")
options = webdriver.ChromeOptions()
options.add_extension('uBlock-Origin.crx')
driver = webdriver.Chrome(service=service, options=options)


driver.maximize_window()
# launch browser
driver.get("https://www.wuxiaworld.eu/chapter/apartment-for-rent-1")


while (1):

    title_xpath = '//*[@id="__next"]/div/div[3]/div[1]/div[1]/div/div[2]/h1'

    titleElementList = driver.find_elements(By.XPATH, title_xpath)

    if (titleElementList != []):
        pass
    else:
        break

    titleEl = driver.find_element(By.XPATH, title_xpath)
    chapter_heading = titleEl.text

    content = driver.find_element(
        By.XPATH, '//*[@id="__next"]/div/div[3]/div[1]/div[2]/div[1]/div')

    next_btn = driver.find_element(
        By.XPATH, '//*[@id="__next"]/div/div[3]/div[1]/div[1]/div/div[4]/a[2]')

    Scrape_Site(content, chapter_heading)

    next_btn.click()
    time.sleep(5)


driver.quit()

doc.save(fileNameMain)
