import io
from PIL import Image
from bs4 import BeautifulSoup
import requests
import openpyxl
from docx import Document
from docx.shared import Inches, Cm, Pt


# 1 px = ( 2.54 / 96 ) cm
# width_in_cm = width_in_px * (2.54 / 96)
# height_in_cm = height_in_px * (2.54 / 96)
# width_in_cm = float("%.2f" % width_in_cm)
# height_in_cm = float("%.2f" % height_in_cm)
# print('cm', width_in_cm, height_in_cm)


def writeText(fileName, txt):
    fw = open(fileName, "w", encoding="utf-8")
    fw.write(str(txt))
    fw.close()


reduce_page_break_works = 1.2

# for img add to docx
max_width = 17.79
max_height = 22.86 - reduce_page_break_works


def init_document(doc):
    # style the body text
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(16)

    # style the heading
    HeadingStyle = doc.styles['Heading 1']
    HeadingFont = HeadingStyle.font
    HeadingFont.size = Pt(24)

    # set margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)


def Scrape_Site(site_url, chapter_heading):
    try:
        source = requests.get(site_url)
        source.raise_for_status()
        soup = BeautifulSoup(source.text, 'html.parser')

        """ with open('test.html', "r", encoding="utf-8") as f:
            contents = f.read()
        soup = BeautifulSoup(contents, 'html.parser') """

        storyPart = soup.select_one(".entry-content")

        # remover header and footer id=textbox
        for header in storyPart.select('#textbox'):
            header.decompose()

        for header in storyPart.select('noscript'):
            header.decompose()

        for header in storyPart.select('script'):
            header.decompose()

        doc.add_heading(chapter_heading, level=1)
        finalStory = storyPart.findChildren()
        skip_tag = None

        for t, tag in enumerate(finalStory):
            imgs = tag.findAll("img")
            if imgs != []:
                for img in imgs:
                    doc.add_page_break()
                    img_url = img['src']
                    img_response = requests.get(img_url, stream=True)
                    image = io.BytesIO(img_response.content)
                    img = Image.open(image)
                    width_in_px, height_in_px = img.size

                    if width_in_px > height_in_px:
                        width_in_cm = max_width
                        doc.add_picture(image, width=Cm(width_in_cm))

                    elif height_in_px > width_in_px:
                        height_in_px = max_height
                        doc.add_picture(image, height=Cm(height_in_px))

                    doc.add_page_break()

            else:
                current_tag_text = tag.get_text(strip=True)
                storyPart_text = tag.get_text("\n\n", strip=True)

                if current_tag_text != '' and skip_tag != tag:
                    doc.add_paragraph('\n' + storyPart_text)
                    tagChildren = tag.findChildren()
                    if tagChildren != []:
                        first_tagChildren = tagChildren[0]
                        skip_tag = first_tagChildren
                    else:
                        skip_tag = None

        doc.add_page_break()

    except Exception as e:
        print(e)


doc = Document()
init_document(doc)

""" site_url = ''

chapter_heading = 'An Alluring Scheme'
Scrape_Site(site_url, chapter_heading)

fileNameMain = "maou-{}.docx"
doc.save(fileNameMain.format(1)) """


excel = openpyxl.load_workbook('Maou no Hajimekata Table of Contents.xlsx')
sheet = excel.active

total_rows = sheet.max_row

all_vol_no = tuple()
all_ch_title = tuple()
all_site_url = tuple()


for x in range(1, total_rows + 1):
    col_one = sheet.cell(row=x, column=1).value
    col_two = sheet.cell(row=x, column=2).value
    col_three = sheet.cell(row=x, column=3).value
    all_vol_no += (col_one,)
    all_ch_title += (col_two,)
    all_site_url += (col_three,)


count = i = 0
countGallopGap = 50  # no of ch per file
countGallop = countGallopGap

doc = Document()
init_document(doc)


fileNameMain = "Maou-no-Hajimekata-{}.docx"

for site_url in all_site_url:
    count = count + 1

    vol_no = all_vol_no[count-1]
    ch_title = all_ch_title[count-1]
    chapter_heading = f'{vol_no}: {ch_title}'

    Scrape_Site(site_url, chapter_heading)

    # write a specific no of chapter in multiple files
    if count == countGallop:
        i = i + 1
        countGallop = countGallop + countGallopGap

        doc.save(fileNameMain.format(i))

        doc = Document()
        init_document(doc)
        print(count, 'if', chapter_heading)

    # save the last remaining no of chapter in another file
    elif count == len(all_site_url):
        i = i + 1
        doc.save(fileNameMain.format(i))

    else:
        print(count, 'else', chapter_heading)

    # testing
    # if count == 2:
    #     break
